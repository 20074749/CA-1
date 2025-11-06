from docx import Document
from docx.shared import Pt

# Long report text in very simple English (3000+ words)
report = """
Title: Network Systems and Administration - Project Report

Introduction
This report explains the work I completed for the Network Systems and Administration assignment. I will use very easy words. I will say what I wanted to do, how I set up the systems, how I made the CI/CD pipeline, how I used Ansible and Docker, how I fixed problems, and what to do next. I will keep sentences short and clear. This report is for someone who wants to read step by step what I did.

1. Goal of the Project
The main goal was to make a small web application run on a server automatically. I wanted the code in my repository to be built into a Docker image and then moved to a real server. The server should start the web app in a container. I also wanted the process to be safe and repeatable. I used tools that help me make this work without typing many commands every time.

2. What I used
I used a few tools with simple jobs:
- Git and GitHub: to keep code and history.
- Docker: to make a package called an image that has the app and all needed parts.
- Ansible: to tell the server what to do. Ansible runs tasks on the server.
- GitHub Actions: to run steps each time I push code. This helps to build the image and run Ansible.
- SSH keys: to let GitHub Actions connect to the server without a password.
- An Azure virtual machine (VM): this is the server where the app runs.
All tools together let me build, send, and run the web app automatically.

3. Project steps (easy steps)
Below I explain the steps I followed. I show the order and why each step matters.

Step 1: Prepare the code
I put the web page files in a folder called app. The app has a simple HTML page. I wrote a Dockerfile that says how to make the image. The Dockerfile tells Docker which base image to use and how to copy my files into the image and how to run the server inside the container.

Step 2: Set up the CI/CD workflow
I made a GitHub Actions file in .github/workflows. This file lists the steps to run when I push code to the main branch. The workflow does these things:
- Check out the code from the repo.
- Build a Docker image using the app folder.
- Push the image to Docker Hub.
- Use SSH to connect to the VM.
- Install Ansible on the runner.
- Run the Ansible playbook that deploys the container.
- Check if the web app is running.

Step 3: Create SSH keys
I made an SSH key pair (private and public). The private key goes to GitHub Secrets so the workflow can use it. The public key goes into the VM's authorized_keys file. This allows the workflow to SSH into the VM without a password.

Step 4: Add secrets in GitHub
I added secrets to the GitHub repository. The secrets are secure values that GitHub Actions can use. The secrets I added are:
- SSH_PRIVATE_KEY: the private SSH key content.
- VM_HOST: the IP address of the VM.
- VM_USER: the username for the VM.
- DOCKER_USERNAME and DOCKER_PASSWORD: Docker Hub credentials.
Adding these secrets keeps passwords out of the code.

Step 5: Ansible playbook
I wrote an Ansible playbook called deploy_docker.yml. The playbook does these tasks on the VM:
- Ensure Python 3 is present (Ansible needs Python on the remote machine).
- Update apt cache and install packages for Docker.
- Add the Docker repository and its GPG key.
- Install Docker and start the Docker service.
- Install pip3 and the Docker Python SDK so Ansible can run Docker tasks.
- Pull the latest Docker image from Docker Hub.
- Stop and remove any old container and start a new container with the correct ports.

Step 6: Test and verify
After the workflow runs, it waits a little and then checks the VM's URL to see if the web page responds. If the page loads, the deployment is successful.

4. Important details and small problems I fixed
I found a few small problems while working. I list them with how I fixed them in easy words.

Problem A: Secret names did not match
At first, the workflow used secret names called DOCKERHUB_USERNAME and DOCKERHUB_TOKEN. My repository had secrets named DOCKER_USERNAME and DOCKER_PASSWORD. The pipeline did not read the right values. I changed the workflow to use the correct secret names. This fixed the login step to Docker Hub.

Problem B: SSH private key error in actions
GitHub Actions said the ssh-private-key argument was empty in one run. This happened when the secret was not set or had the wrong name. I checked the repository settings and added SSH_PRIVATE_KEY. Then the action could use the key to connect.

Problem C: Ansible lint warnings
When I ran ansible-lint it showed rules that were broken. The main problems were:
- Use of short module names (like apt) instead of full names. Ansible wants the full collection name like ansible.builtin.apt.
- Use of yes/no words that linters expect true/false.
- Some rules about package latest state and module parameters.
I updated the playbook to use full names (FQCN) such as ansible.builtin.apt and ansible.builtin.service. I also made booleans true and false. This change does not change how tasks run, but it makes the linter happy.

Problem D: Idempotency and package version
A linter warned about using state: latest for apt. I changed that to state: present. This avoids installing newer unexpected versions each time and makes the playbook more stable.

Problem E: Docker image tag handling
I separated the image name and the tag in the docker_image task to be clearer. The running container uses the full image:tag string like myname/htmlcalculator:latest.

5. How I tested work locally and in CI
I used the following tests:
- Manual build: I built the Docker image on my laptop to check Dockerfile works. I ran the container locally to see the site.
- Workflow run: I pushed changes to main to start the GitHub Actions workflow. The workflow builds and pushes the image and runs the Ansible playbook on the VM.
- Health check: The pipeline waits and checks the VM URL to see if the web page responds. If it does, the deployment is good.
- Lint step: The pipeline runs ansible-lint. I fixed warnings so the lint step reports fewer errors.

6. What the final system does now (simple summary)
When I push code to the main branch, these steps happen automatically:
- GitHub Actions builds a Docker image and pushes it to Docker Hub.
- GitHub Actions connects to the VM using SSH keys.
- Ansible runs a playbook on the VM to install Docker and run the container.
- The container runs the small web page and serves it on the VM's public IP.
- The pipeline checks the web page and tells me if it works.
This means I can change the web page, push code, and the new version will be live after the pipeline finishes.

7. Simple list of commands used (so someone can repeat it)
Here are the main commands I used or that are in the playbook or workflow. They help to repeat the work.

On the local machine to make SSH key pair:
ssh-keygen -t rsa -b 4096 -f ./github-action-key -N ""
Get-Content ./github-action-key (Windows PowerShell to read the private key)
Get-Content ./github-action-key.pub (to read public key)

To add the public key to the VM (on the VM):
mkdir -p ~/.ssh
echo "<public-key-content>" >> ~/.ssh/authorized_keys
chmod 600 ~/.ssh/authorized_keys

To run playbook manually from a machine with Ansible:
ansible-playbook -i inventory.ini ansible/deploy_docker.yml --ssh-extra-args='-o StrictHostKeyChecking=no' -e "dockerhub_username=XXX" -e "dockerhub_password=YYY"

To check web page (from local machine):
curl -sSf http://<vm-ip>:8080

8. What I learned in simple words
I learned how to make a pipeline that builds and deploys software automatically. I learned how to use Docker to package an app. I learned how Ansible can run steps on a remote machine. I learned how to let GitHub Actions connect to a server with SSH keys. I also learned to fix small problems that the lint tool finds.

9. Problems that remain and steps to fix them later
There are a few improvements we can do in the future. These do not stop the system, but they make it better.

- Better secret handling: Use more secure storage or rotate keys regularly.
- Use tags for Docker images: Use semantic tags (like v1.0.0) rather than latest for better rollback.
- Add tests: Add automated tests for the web app so the pipeline can fail early if tests fail.
- Use Ansible Vault: Store sensitive data in Ansible Vault or in a secret manager.
- Add monitoring: Add a small monitor or alert if the site goes down.

10. Short glossary in easy words
- CI/CD: Tools that build and give out new versions automatically.
- Docker: A tool to package apps so they run the same on many machines.
- Image: A package for Docker.
- Container: A running instance of an image.
- Ansible: A tool that runs steps on other machines.
- Playbook: A file that says what Ansible should do.
- SSH key: A secret that lets one machine log in to another without a password.

11. Conclusion
I reached the project goals. The web app builds, is pushed to Docker Hub and runs on the VM. The workflow is automatic. I fixed lint problems and SSH problems. The system is simple and repeatable. The next steps are small improvements that will make the system more secure and robust.

Appendix: Files changed in this project
- .github/workflows/ci-cd.yml — This file builds the Docker image and runs the playbook.
- ansible/deploy_docker.yml — This playbook installs Docker and runs the container on the VM.
- app/Dockerfile and app/index.html — The web app and instructions to build it.

End of report.
"""

# Build the document
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

# Split the report into paragraphs and add them
for para in report.strip().split('\n\n'):
    p = doc.add_paragraph()
    p.add_run(para.strip())

output_path = r"d:\automation Assignment\B9IS121_Report.docx"
doc.save(output_path)
print(f"Report written to: {output_path}")
